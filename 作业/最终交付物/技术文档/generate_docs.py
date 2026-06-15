"""Generate 4 Word docs for 虞舍 + 数字苔藓 project."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.dirname(__file__)

def make_doc(title, doc_id, sections):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # Header style
    for lvl in [1, 2, 3]:
        hs = doc.styles[f'Heading {lvl}']
        hs.font.name = 'Microsoft YaHei'
        hs.font.color.rgb = RGBColor(0x3D, 0x30, 0x27)

    # Title page
    doc.add_paragraph('')
    doc.add_paragraph('')
    t = doc.add_paragraph(title)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(22)
    t.runs[0].font.bold = True
    t.runs[0].font.color.rgb = RGBColor(0x3D, 0x30, 0x27)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'\n文档编号：{doc_id}\n版本：V2.0\n编写：实训第4组\n日期：2026-07-05\n').font.size = Pt(11)

    doc.add_page_break()

    for sec in sections:
        if sec.get('title'):
            doc.add_heading(sec['title'], level=1)
        if sec.get('text'):
            doc.add_paragraph(sec['text'])
        if sec.get('items'):
            for item in sec['items']:
                doc.add_paragraph(item, style='List Bullet')
        if sec.get('subs'):
            for sub in sec['subs']:
                doc.add_heading(sub['title'], level=2)
                if sub.get('text'):
                    doc.add_paragraph(sub['text'])
                if sub.get('items'):
                    for item in sub['items']:
                        doc.add_paragraph(item, style='List Bullet')

    path = os.path.join(OUT, f'{os.path.basename(title)}.docx')
    doc.save(path)
    return path

# ================================================================
# DOC 1: 项目需求分析
# ================================================================
s1 = [
    {'title': '1. 项目背景', 'text': (
        '虞舍是一个个人独立博客项目，部署在校内服务器（10.42.78.75），对外提供文章阅读、评论互动、RSS订阅等服务。'
        '项目采用单文件全栈架构——前端为纯 HTML+CSS+JS 编写的单文件博客，后端为 Python 单文件 HTTP Server，'
        '数据以 JSON 文件形式持久化存储。\n\n'
        '数字苔藓是虞舍博客的衍生项目，一个活的网页花园。通过 Canvas 2D 程序化渲染，实现了 6 种植物在四季昼夜环境下的'
        '生长、枯萎、浇水互动等完整生态模拟。每次页面访问都会自动播种新植物，访客可以点击浇水照料植物，'
        '植物随时间自然衰减健康值，最终枯萎或开花。'
    )},
    {'title': '2. 功能需求', 'subs': [
        {'title': '2.1 博客功能模块', 'items': [
            '文章系统：支持 Markdown 格式文章，分类管理（技术/生活/随笔/分享），标签系统，封面图，置顶',
            '阅读体验：全文搜索（按 / 键聚焦），时间线归档，热门文章 TOP 10，Canvas 鸟瞰图可视化',
            '互动系统：文章评论，RSS 订阅（/api/blog/rss），友链页面，随机阅读（🎲），代码块一键复制',
            '特殊功能：时间胶囊（定时发布），文章海报生成下载（Canvas 渲染），暗黑模式切换',
            '管理面板：Ctrl+K 切换管理员模式，发布/编辑/删除文章，置顶管理（密码保护）',
        ]},
        {'title': '2.2 数字苔藓功能模块', 'items': [
            '植物系统：6 种植物类型（蕨/花/蘑菇/草/多肉/藤蔓），每种 3 个形态变体，程序化绘制的生长动画',
            '生命周期：种子→嫩芽→成长→开花→枯萎，健康值 0-100，自然衰减 -0.5/小时，浇水 +20',
            '浇水互动：点击/触摸植物区域，粒子雨滴动画，植物弹跳反馈，每分钟限浇 3 次',
            '季节系统：春/夏/秋/冬四季，影响植物生成权重和生长速度，冬季飘雪',
            '昼夜循环：根据真实时间渐变天空色温，夜晚萤火虫更亮更密，月亮和星星出现',
            '生物生态：蜗牛沿土丘爬行+黏液轨迹，蝴蝶随机飞行，萤火虫漂浮闪烁',
            '木牌语录：10 条轮播语录，双击木牌晃动落树叶并切换',
        ]},
    ]},
    {'title': '3. 非功能需求', 'items': [
        '性能：博客页面加载时间 < 2 秒，花园 Canvas 渲染保持 60fps',
        '可靠性：JSON 文件持久化，数据即时写入，异常情况下不丢失超过 1 次操作的数据',
        '安全性：管理面板密码保护，API 端点权限验证，XSS 防护',
        '兼容性：支持 Chrome/Firefox/Edge 主流浏览器，移动端响应式适配',
        '可维护性：单文件架构，零外部依赖，部署仅需 Python 3.6+',
    ]},
    {'title': '4. 技术约束', 'items': [
        '部署环境：校内服务器 (CentOS)，内网 IP 10.42.78.75',
        '运行环境：Python 3.9+ 自带的 http.server 模块',
        '前端约束：纯 HTML+CSS+JS，不使用任何前端框架或构建工具',
        '存储约束：JSON 文件存储，不使用数据库',
        '网络约束：80 端口通过 Nginx 反向代理对外服务，8089 端口仅内网可访问',
    ]},
]

# ================================================================
# DOC 2: 系统架构设计
# ================================================================
s2 = [
    {'title': '1. 总体架构', 'text': (
        '虞舍+数字苔藓采用单文件全栈架构，分为四层：\n\n'
        '┌─────────────────────────────────┐\n'
        '│  用户浏览器 (Chrome/Firefox/Edge) │\n'
        '├─────────────────────────────────┤\n'
        '│  Nginx :80 (反向代理+静态文件)      │\n'
        '│  / → yushe-blog.html              │\n'
        '│  /moss → proxy → wall_api.py      │\n'
        '│  /garden/ → proxy → wall_api.py   │\n'
        '├─────────────────────────────────┤\n'
        '│  wall_api.py :8089 (Python HTTP)  │\n'
        '│  GET/POST/PUT/DELETE 路由器        │\n'
        '├─────────────────────────────────┤\n'
        '│  JSON 文件存储                     │\n'
        '│  blog_articles.json / garden.json │\n'
        '└─────────────────────────────────┘'
    )},
    {'title': '2. 前端架构', 'subs': [
        {'title': '2.1 博客前端 (yushe-blog.html)', 'items': [
            '单文件 HTML，约 635 行，包含完整的 CSS + JS',
            'SPA 模式：所有页面切换通过 JS 渲染，无页面刷新',
            '视图系统：文章列表/文章详情/归档/标签云/友链/热榜/鸟瞰/时间胶囊',
            'Canvas 子模块：鸟瞰图（文章关系可视化）、海报生成',
            '状态管理：全局 allArticles 数组 + sessionStorage 管理登录/暗黑',
        ]},
        {'title': '2.2 花园前端 (moss.html)', 'items': [
            '单文件 HTML，约 642 行 Canvas 2D 渲染',
            '18 层渲染管线：天空→星星→月亮→光照→远树→土丘→卵石→落叶→野菇→露珠→雾→雪→雨→萤火虫→微粒→蝴蝶→蜗牛→植物→暗角',
            '程序化植物绘制：6 种植物各有 3 个变体，通过贝塞尔曲线/椭圆/弧线组合绘制',
            '状态管理：全局 S 对象管理植物列表、季节、粒子、动画',
            'API 客户端：X-Session-Id 头部标识会话，60 秒轮询同步花园状态',
        ]},
    ]},
    {'title': '3. 后端架构 (wall_api.py)', 'subs': [
        {'title': '3.1 技术选型', 'items': [
            'Python http.server.BaseHTTPRequestHandler — 标准库，零依赖',
            'JSON 文件存储 — 读写简单，备份=复制文件',
            '单线程处理 — 无并发竞争，适合低流量场景',
        ]},
        {'title': '3.2 API 设计', 'items': [
            '博客 API：GET/POST/PUT/DELETE /blog/articles — 文章 CRUD',
            '评论 API：POST /blog/articles/:id/comments — 发表评论',
            '阅读计数：POST /blog/articles/:id/view — 浏览量+1',
            'RSS：GET /blog/rss — XML 格式 RSS Feed',
            '花园状态：GET /garden/state — 全部植物+元信息',
            '访问播种：POST /garden/visit — 创建新植物，sessionStorage 去重',
            '浇水：POST /garden/water — 距离检测+健康更新+限流',
        ]},
        {'title': '3.3 数据模型', 'items': [
            'Plant: {id, type, x, y, size, health, stage, variant, createdAt, lastWateredAt, wateredBy[]}',
            'Garden: {totalVisits, lastVisitAt, dailyWaterings{}, plants[]}',
            '容量控制：上限 200 株植物，超出时优先移除枯萎>非开花>最老',
            '衰减系统：每次读取时根据 lastWateredAt 计算健康衰减',
        ]},
    ]},
    {'title': '4. 部署架构', 'items': [
        '校内服务器 (10.42.78.75)，CentOS 操作系统',
        'Nginx 1.25 作为前端反向代理，监听 80 端口',
        'wall_api.py 通过 nohup 后台运行在 8089 端口',
        '数据文件存储在 /opt/blog/data/ 目录',
        '部署方式：scp 上传文件 → ssh 重启进程',
    ]},
]

# ================================================================
# DOC 3: 部署实施手册
# ================================================================
s3 = [
    {'title': '1. 环境要求', 'items': [
        '操作系统：CentOS 7+ / Ubuntu 18+ 或任何支持 Python 3.6+ 的 Linux 发行版',
        'Python 版本：3.6 及以上（仅需标准库，无需 pip install）',
        'Web 服务器：Nginx 1.20+（用于反向代理和静态文件服务）',
        '网络：内网 IP 10.42.78.75，开放 80 端口',
        '磁盘空间：至少 100MB（JSON 数据文件 + HTML 文件 + 素材）',
    ]},
    {'title': '2. 部署步骤', 'subs': [
        {'title': '2.1 上传文件', 'text': (
            '将以下文件上传到服务器的 /opt/blog/nginx/ 目录：\n'
            '- yushe-blog.html（博客前端）\n'
            '- moss.html（数字苔藓花园）\n'
            '- wall_api.py（后端 API 服务）\n'
            '- moss-assets/（花园素材，可选）'
        )},
        {'title': '2.2 配置 Nginx', 'text': (
            '在 nginx.conf 中添加以下 location 块：\n\n'
            '# 博客首页\n'
            'location = / { root /opt/blog/nginx; try_files /yushe-blog.html =404; }\n\n'
            '# 数字苔藓花园\n'
            'location /moss { proxy_pass http://127.0.0.1:8089/moss; }\n\n'
            '# 花园 API\n'
            'location /garden/ { proxy_pass http://127.0.0.1:8089/garden/; }\n\n'
            '# 博客 API\n'
            'location /api/ { proxy_pass http://127.0.0.1:8089/; }\n\n'
            '# 花园素材\n'
            'location /moss-assets/ { proxy_pass http://127.0.0.1:8089/moss-assets/; }'
        )},
        {'title': '2.3 启动后端服务', 'text': (
            '使用 nohup 在后台启动 wall_api.py：\n\n'
            '$ cd /opt/blog/nginx\n'
            '$ nohup python3 wall_api.py > /dev/null 2>&1 &\n\n'
            '验证启动：\n'
            '$ curl http://localhost:8089/garden/state\n'
            '{"totalVisits": 0, "plants": [], "season": "summer", "month": 6}'
        )},
        {'title': '2.4 重载 Nginx', 'text': (
            '更新配置后重载 Nginx：\n\n'
            '$ nginx -s reload\n'
            '或使用容器：\n'
            '$ kill -HUP $(pgrep -f "nginx.*master")'
        )},
    ]},
    {'title': '3. 访问验证', 'items': [
        '博客首页：http://10.42.78.75',
        '数字苔藓：http://10.42.78.75/moss',
        '花园 API：http://10.42.78.75/garden/state',
        '博客 RSS：http://10.42.78.75/api/blog/rss',
    ]},
    {'title': '4. 快速部署命令', 'text': (
        '# 一键部署脚本（从本地项目目录执行）：\n\n'
        'scp yushe-blog.html moss.html wall_api.py root@10.42.78.75:/opt/blog/nginx/\n'
        'scp -r moss-assets/ root@10.42.78.75:/opt/blog/nginx/\n'
        'ssh root@10.42.78.75 "kill \\$(pgrep -f wall_api.py); cd /opt/blog/nginx && nohup python3 wall_api.py &>/dev/null &"\n'
        'ssh root@10.42.78.75 "kill -HUP \\$(pgrep -f \'nginx.*master\')"'
    )},
    {'title': '5. 故障排查', 'items': [
        '后端未启动：curl http://localhost:8089/garden/state 检查是否返回 JSON',
        '端口被占用：fuser -k 8089/tcp 杀掉旧进程后重启',
        'Nginx 报 502：确认 wall_api.py 在运行，并监听 8089 端口',
        '页面 404：检查 Nginx location 配置和文件路径',
        '素材不显示：确认 moss-assets/ 目录存在于 /opt/blog/nginx/ 下',
    ]},
]

# ================================================================
# DOC 4: 运维操作手册
# ================================================================
s4 = [
    {'title': '1. 日常巡检', 'items': [
        '检查后端进程：ps aux | grep wall_api.py',
        '检查 API 可用性：curl http://localhost:8089/garden/state',
        '检查 Nginx 状态：curl -I http://localhost',
        '检查磁盘空间：df -h /opt/blog/data',
        '检查数据文件：ls -la /opt/blog/data/*.json',
    ]},
    {'title': '2. 数据备份', 'subs': [
        {'title': '2.1 手动备份', 'text': (
            '直接复制 JSON 数据文件：\n\n'
            '$ cp /opt/blog/data/blog_articles.json /opt/blog/data/blog_articles.json.bak.$(date +%Y%m%d)\n'
            '$ cp /opt/blog/data/garden.json /opt/blog/data/garden.json.bak.$(date +%Y%m%d)'
        )},
        {'title': '2.2 自动备份 (crontab)', 'text': (
            '添加 crontab 定时任务，每日凌晨 2 点自动备份：\n\n'
            '$ crontab -e\n'
            '0 2 * * * cp /opt/blog/data/blog_articles.json /opt/blog/backups/blog_$(date +\\%Y\\%m\\%d).json\n'
            '0 2 * * * cp /opt/blog/data/garden.json /opt/blog/backups/garden_$(date +\\%Y\\%m\\%d).json'
        )},
        {'title': '2.3 数据恢复', 'text': (
            '将备份文件复制回原位置，重启后端：\n\n'
            '$ cp /opt/blog/backups/blog_20260701.json /opt/blog/data/blog_articles.json\n'
            '$ cp /opt/blog/backups/garden_20260701.json /opt/blog/data/garden.json\n'
            '$ kill $(pgrep -f wall_api.py); cd /opt/blog/nginx && nohup python3 wall_api.py &'
        )},
    ]},
    {'title': '3. 进程管理', 'items': [
        '启动：cd /opt/blog/nginx && nohup python3 wall_api.py > /dev/null 2>&1 &',
        '停止：kill $(pgrep -f wall_api.py)',
        '重启：kill $(pgrep -f wall_api.py); sleep 1; cd /opt/blog/nginx && nohup python3 wall_api.py > /dev/null 2>&1 &',
        '查看日志：tail -f /var/log/nginx/error.log',
        '检查端口：netstat -tlnp | grep 8089',
    ]},
    {'title': '4. 更新部署', 'text': (
        '从本地推送新版本到服务器：\n\n'
        'scp yushe-blog.html moss.html wall_api.py root@10.42.78.75:/opt/blog/nginx/\n'
        'ssh root@10.42.78.75 "kill \\$(pgrep -f wall_api.py); cd /opt/blog/nginx && nohup python3 wall_api.py &>/dev/null &"\n'
        'ssh root@10.42.78.75 "kill -HUP \\$(pgrep -f \'nginx.*master\')"\n\n'
        '更新素材：\n'
        'scp -r moss-assets/*.png root@10.42.78.75:/opt/blog/nginx/moss-assets/'
    )},
    {'title': '5. 监控指标', 'items': [
        'API 响应时间：正常情况下 < 50ms（JSON 文件读写）',
        '花园植物数量：GET /garden/state 查看 plants 数组长度',
        '博客文章数：GET /api/blog/articles 查看数组长度',
        '磁盘使用：JSON 文件通常 < 10MB，定期检查',
        '内存使用：Python 进程通常 < 50MB',
    ]},
    {'title': '6. 常见问题', 'items': [
        'Q: 博客文章丢失？A: 检查 /opt/blog/data/blog_articles.json 是否存在且格式正确，从备份恢复',
        'Q: 花园植物全部枯萎？A: 正常现象——植物随时间衰减，浇水可以恢复；自然衰减速率 -0.5/小时',
        'Q: 素材图片不显示？A: 检查 moss-assets/ 目录中是否有对应的 PNG 文件，花园会自动降级为程序化绘制',
        'Q: API 返回 404？A: 检查 Nginx 配置中 /api/ 和 /garden/ 的 proxy_pass 是否正确',
        'Q: 服务器重启后服务未自动启动？A: 需要手动执行 nohup 命令重启 wall_api.py',
    ]},
]

# ================================================================
# Generate all 4 docs
# ================================================================
docs = [
    ('1-项目需求分析', 'DMP-REQ-2026-001', s1),
    ('2-系统架构设计', 'DMP-ARCH-2026-001', s2),
    ('3-部署实施手册', 'DMP-DEPLOY-2026-001', s3),
    ('4-运维操作手册', 'DMP-OPS-2026-001', s4),
]

for name, doc_id, sections in docs:
    path = make_doc(name, doc_id, sections)
    print(f'Generated: {path}')
